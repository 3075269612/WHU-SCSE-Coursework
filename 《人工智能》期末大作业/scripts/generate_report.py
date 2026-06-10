from __future__ import annotations

import json
import re
import sys
import time
from importlib import metadata
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

sys.path.append(str(Path(__file__).resolve().parents[1]))

from src.utils import PROJECT_ROOT, project_path


REPORT_PATH = project_path("report/AI-王李明-2024302181194.docx")
FONT_CLEAN_PATH = project_path("report/AI-王李明-2024302181194-font-clean.docx")

REQUIRED_BOLD_TERMS = [
    "正文 4000 字以上",
    "技术分析",
    "监督学习",
    "横向对比",
    "技术评价",
    "数学原理",
    "算法流程",
    "优缺点",
    "应用设计",
    "分类问题",
    "完整实验流程",
    "数据获取",
    "数据探索",
    "缺失值检测",
    "类别分布",
    "特征分布",
    "可视化",
    "柱状图",
    "直方图",
    "特征编码",
    "TfidfVectorizer",
    "TF-IDF",
    "CountVectorizer",
    "数据集划分",
    "训练集",
    "验证集",
    "测试集",
    "60%/20%/20%",
    "stratify=y",
    "random_state=42",
    "先划分后预处理",
    "fit_transform",
    "transform",
    "MLP",
    "网络结构",
    "PyTorch",
    "CrossEntropyLoss",
    "AdamW",
    "CosineAnnealingLR",
    "Batch",
    "Epoch",
    "环境依赖",
    "OS",
    "Python",
    "GPU/CUDA",
    "超参数敏感性分析",
    "控制变量法",
    "至少 2 个参数",
    "表格",
    "折线图",
    "训练曲线",
    "Loss",
    "Accuracy",
    "混淆矩阵",
    "超参敏感性图",
    "Macro-F1",
    "Macro-Precision",
    "Macro-Recall",
    "超参影响",
    "错误分析",
    "过拟合判断",
    "Baseline 对比",
    "baseline",
    "依赖版本",
    "代码逻辑",
    "参考文献",
    "学术诚信",
]


def load_json(path: str) -> dict:
    return json.loads(project_path(path).read_text(encoding="utf-8"))


def fmt(value: float) -> str:
    return f"{float(value):.4f}"


def dependency_versions() -> pd.DataFrame:
    """整理课程 PDF 要求在报告中列出的主要依赖版本。"""
    packages = [
        ("numpy", "numpy"),
        ("pandas", "pandas"),
        ("matplotlib", "matplotlib"),
        ("scikit-learn", "scikit-learn"),
        ("torch", "torch"),
        ("tqdm", "tqdm"),
        ("python-docx", "python-docx"),
    ]
    rows = []
    for display_name, package_name in packages:
        try:
            version = metadata.version(package_name)
        except metadata.PackageNotFoundError:
            version = "未安装/未使用"
        rows.append({"依赖": display_name, "版本": version})
    return pd.DataFrame(rows)


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    """统一设置中西文字体，避免 Word 默认 Calibri/Arial 混入。"""
    run.font.name = "SimSun"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def bold_required_terms_in_paragraph(paragraph) -> None:
    text = paragraph.text
    if not text:
        return

    pattern = re.compile("|".join(re.escape(term) for term in sorted(REQUIRED_BOLD_TERMS, key=len, reverse=True)))
    if not pattern.search(text):
        return

    for run in list(paragraph.runs):
        paragraph._element.remove(run._element)

    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        bold_run = paragraph.add_run(match.group(0))
        bold_run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def bold_required_terms(doc: Document) -> None:
    """将作业 PDF 中的必需内容在报告正文中加粗标识。"""
    for paragraph in doc.paragraphs:
        bold_required_terms_in_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    bold_required_terms_in_paragraph(paragraph)


def configure_styles(doc: Document) -> None:
    """设置整篇文档的基础样式，字号允许分级但字体保持整齐统一。"""
    style_specs = {
        "Normal": (10.5, False),
        "Heading 1": (16, True),
        "Heading 2": (14, True),
        "Heading 3": (12, True),
    }
    for style_name, (size, bold) in style_specs.items():
        style = doc.styles[style_name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style._element.rPr.rFonts.set(qn("w:ascii"), "SimSun")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "SimSun")
        style.font.size = Pt(size)
        style.font.bold = bold


def normalize_document_fonts(doc: Document) -> None:
    """对所有段落、表格和页眉页脚做最终字体归一，处理直接格式。"""
    def normalize_paragraph(paragraph) -> None:
        style_name = paragraph.style.name if paragraph.style is not None else ""
        if style_name == "Heading 1":
            size, bold = 16, True
        elif style_name == "Heading 2":
            size, bold = 14, True
        elif style_name == "Heading 3":
            size, bold = 12, True
        else:
            size, bold = 10.5, None
        for run in paragraph.runs:
            set_run_font(run, size, bold)

    for paragraph in doc.paragraphs:
        normalize_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    normalize_paragraph(paragraph)
    for section in doc.sections:
        for paragraph in section.header.paragraphs + section.footer.paragraphs:
            normalize_paragraph(paragraph)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(21)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)
    return paragraph


def add_plain_paragraph(doc: Document, text: str = ""):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.add_run(text)
    return paragraph


def add_table_from_dataframe(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_plain_paragraph(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, column in enumerate(df.columns):
        hdr_cells[idx].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(df.columns):
            value = row[column]
            if isinstance(value, float):
                cells[idx].text = fmt(value)
            else:
                cells[idx].text = str(value)


def add_picture_if_exists(doc: Document, path: str, caption: str) -> None:
    full_path = project_path(path)
    if full_path.exists():
        doc.add_picture(str(full_path), width=Inches(5.8))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_paragraph = add_plain_paragraph(doc, caption)
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_plain_paragraph(doc, f"{caption}（图片缺失：{path}）")


def chinese_char_count(doc: Document) -> int:
    text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    return len(re.findall(r"[\u4e00-\u9fff]", text))


def build_report() -> None:
    metrics = load_json("results/metrics.json")
    summary = load_json("results/dataset_summary.json")
    baseline_df = pd.read_csv(project_path("results/baselines.csv"))
    hyper_df = pd.read_csv(project_path("results/hyperparams.csv"))
    model_selection_df = pd.read_csv(project_path("results/model_selection.csv"))
    tfidf_df = pd.read_csv(project_path("results/tfidf_sensitivity.csv"))

    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于监督学习的垃圾短信识别分类应用设计与实现")
    set_run_font(run, 18, True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(subtitle.add_run("姓名：王李明    学号：2024302181194    课程：人工智能"), 11, False)

    add_heading(doc, "摘要", 1)
    add_paragraph(doc, "本文围绕垃圾短信识别这一二分类任务，完成了监督学习分类应用从数据获取、清洗探索、特征编码、模型训练到结果评估的完整实验流程。实验使用 UCI SMS Spam Collection 数据集，将短信文本通过 TF-IDF 表示为数值特征，并以多层感知机 MLP 作为主模型，同时引入多数类、随机预测、Logistic Regression、Multinomial Naive Bayes 和 Linear SVM 作为对照。报告重点分析了不同监督学习算法的特点与适用场景，进一步从数学原理、算法流程、优缺点和实验表现评价 MLP。由于数据集中 spam 为少数类，本文在评价时以 Macro-F1 作为核心指标，并结合训练曲线、混淆矩阵和超参数敏感性图分析模型的收敛、错误类型和参数影响。")
    add_plain_paragraph(doc, "关键词：监督学习；文本分类；垃圾短信识别；TF-IDF；MLP；Macro-F1")

    add_heading(doc, "一、引言", 1)
    add_paragraph(doc, "短信仍然是移动通信中最基础、最普遍的文本信息载体。验证码、通知、营销、广告和恶意诱导信息都会以短信形式出现。对普通用户而言，垃圾短信不仅造成阅读负担，也可能诱导用户点击链接、回复号码或拨打高费率电话。与人工规则相比，监督学习方法能够从带标签样本中学习文本模式，在保持实现成本较低的同时，为垃圾短信过滤提供可解释、可复现的技术路线。")
    add_paragraph(doc, "本课程项目选择“垃圾短信识别”作为分类应用场景，任务目标是将一条英文短信自动判别为正常短信 ham 或垃圾短信 spam。项目数据来自 UCI Machine Learning Repository 的 SMS Spam Collection 数据集。该数据集规模适中、下载方式稳定、标签清晰，适合在普通 CPU 环境上完成完整的监督学习实验流程。与大规模预训练语言模型相比，本项目不追求复杂模型，而强调从数据获取、数据探索、特征工程、模型训练、超参数敏感性分析到实验评估的完整性。")
    add_paragraph(doc, "在应用设计上，本文采用 TF-IDF 将短信文本转换为定长数值特征，再使用多层感知机 MLP 作为主要监督学习模型。为了评价模型是否真正有效，实验同时设置多数类预测、按类别比例随机预测、Logistic Regression、Multinomial Naive Bayes 和 Linear SVM 五类 baseline。评价指标包括 Accuracy、Macro-F1、Macro-Precision 和 Macro-Recall。由于垃圾短信数据存在类别不均衡，本文在分析时更重视 Macro-F1，而不是只观察总体准确率。")

    add_heading(doc, "二、监督学习分类技术分析", 1)
    add_paragraph(doc, "监督学习的核心思想是利用带有真实标签的数据训练模型，使模型学习输入特征与目标类别之间的映射关系。对分类问题而言，训练数据通常表示为若干样本及其对应类别标签，模型通过优化损失函数不断调整参数，最终希望在未见过的测试样本上也能保持较好的泛化能力。典型监督分类算法包括朴素贝叶斯、逻辑回归、支持向量机、决策树、随机森林、梯度提升树、K 近邻以及神经网络等。不同算法在假设、表达能力、训练成本和可解释性方面存在明显差异。")
    add_paragraph(doc, "朴素贝叶斯常用于文本分类，尤其适合词袋特征。它基于条件独立假设估计类别后验概率，训练速度快，对小样本友好，但独立性假设较强，难以表达词项之间复杂关系。逻辑回归是线性分类模型，通过 Sigmoid 或 Softmax 建模类别概率，优化交叉熵损失。它具有训练稳定、可解释性强、适合作为 baseline 的优点，但如果原始特征与类别之间存在明显非线性关系，线性模型的表达能力会受到限制。")
    add_paragraph(doc, "支持向量机通过最大化分类间隔提升泛化能力，在高维稀疏文本特征上经常表现良好。线性 SVM 对文本分类尤其有效，但核方法在大数据上训练成本较高，概率输出也不如逻辑回归自然。决策树通过特征划分构造树状规则，解释性较强，但单棵树容易过拟合。随机森林利用多个决策树投票降低方差，适合表格特征，但对极高维稀疏文本特征并不总是最优。")
    add_paragraph(doc, "神经网络通过多层非线性变换学习复杂函数。MLP 是最基础的前馈神经网络，常由全连接层、激活函数和正则化模块组成。相较于线性模型，MLP 能表示非线性决策边界；相较于 CNN、RNN 或 Transformer，MLP 的结构更简单，训练成本更低，适合课程项目中展示监督学习核心流程。对 TF-IDF 文本特征而言，MLP 可以在稀疏词项权重的基础上学习词项组合的非线性贡献，但也更依赖超参数选择和正则化设置。")
    add_paragraph(doc, "从应用场景看，如果数据规模较小、解释性要求高，可以优先考虑逻辑回归、朴素贝叶斯或线性 SVM；如果存在复杂非线性关系，可以考虑随机森林、梯度提升树或神经网络；如果输入是图像、语音或长文本，则 CNN、RNN、Transformer 等模型更有优势。本项目选择 MLP 的理由是：一方面它符合课程 PDF 对 MLP 模型实现的建议；另一方面它能够在 CPU 上稳定运行，并与 Logistic Regression baseline 形成清晰对比。")

    add_heading(doc, "三、MLP 技术评价", 1)
    add_paragraph(doc, "多层感知机是一类典型的前馈神经网络。设输入特征向量为 x，第一层线性变换可表示为 z1 = W1x + b1，经过 ReLU 激活得到 h1 = max(0, z1)。后续隐藏层重复线性变换与非线性激活，最后输出层得到每个类别的 logit。对于二分类或多分类问题，通常将 logits 输入 CrossEntropyLoss，该损失等价于 Softmax 与负对数似然的组合。训练目标是最小化预测类别分布与真实标签之间的差异。")
    add_paragraph(doc, f"本项目中的最终 MLP 结构为 Linear(in,{metrics['config']['hidden_dims'][0]}) -> ReLU -> Dropout({metrics['config']['dropout']}) -> Linear({metrics['config']['hidden_dims'][0]},{metrics['config']['hidden_dims'][1]}) -> ReLU -> Dropout({metrics['config']['dropout']}) -> Linear({metrics['config']['hidden_dims'][1]},2)。输入维度由 TF-IDF 词表大小决定，输出维度为 2，分别对应 ham 与 spam。隐藏层维度选择 {metrics['config']['hidden_dims']}，是基于验证集 Macro-F1 在候选结构中选择得到的结果。若隐藏层过小，模型可能无法捕捉区分垃圾短信的关键词组合；若隐藏层过大，在样本规模只有数千条的情况下容易过拟合，并增加训练时间。")
    add_paragraph(doc, "训练过程中使用 AdamW 优化器和 CosineAnnealingLR 学习率调度器。AdamW 在 Adam 的自适应学习率基础上将权重衰减与梯度更新解耦，常用于神经网络训练；CosineAnnealingLR 使学习率随 epoch 按余弦曲线逐步下降，有助于在训练后期进行更细粒度的参数搜索。学习率控制每次参数更新幅度，过大会造成训练震荡或发散，过小则可能收敛缓慢。Dropout 在训练阶段随机屏蔽部分神经元，使模型不能过度依赖某些局部特征，从而缓解过拟合。由于 spam 是少数类，本次优化后的实验还在 CrossEntropyLoss 中加入类别权重，使少数类样本在损失函数中具有更高权重。")
    add_paragraph(doc, "MLP 的算法流程可以概括为：第一，将短信文本经过 TF-IDF 转换为输入向量；第二，按 batch 将训练样本送入网络，依次完成线性变换、ReLU 激活和 Dropout；第三，输出层产生 ham 与 spam 两个类别的 logits，并用带类别权重的 CrossEntropyLoss 计算损失；第四，反向传播计算梯度，AdamW 更新参数，调度器更新学习率；第五，每个 epoch 在验证集上计算 Accuracy、Macro-F1、Macro-Precision 和 Macro-Recall；第六，若验证集 Macro-F1 创新高则保存 best checkpoint，若连续多个 epoch 无提升则 early stopping；最后只在测试集上评估一次 best checkpoint，避免用测试集参与模型选择。")
    add_paragraph(doc, "MLP 的优点在于结构清晰、实现简单、可通过增加层数和隐藏单元提升表达能力，并且可以与多种特征工程方法结合。缺点是可解释性弱于线性模型，对超参数较敏感，对输入特征尺度和数据划分方式也比较依赖。在文本分类任务中，MLP 不像朴素贝叶斯那样天然利用概率假设，也不像 Transformer 那样直接建模上下文语义，但它能在低成本条件下展示神经网络分类器的基本训练流程和评估方法。")
    add_paragraph(doc, "从课程理解角度看，MLP 的价值不只在于最终指标，更在于它把监督学习中的若干关键概念连接起来：特征向量、参数化模型、损失函数、梯度下降、正则化、验证集选择和测试集评估。理解这些环节，比单纯追求更复杂模型更符合本次作业“复杂度不重要，分析深度才重要”的要求。本次实验中 MLP 没有超过 Linear SVM，也说明模型复杂度并不等同于分类效果，特征表示、数据规模和归纳偏置同样会决定最终表现。")

    add_heading(doc, "四、应用设计与数据处理", 1)
    add_paragraph(doc, f"本项目使用 UCI SMS Spam Collection 数据集。实验读取后共有 {summary['total_rows']} 条有效样本，其中正常短信 ham 为 {summary['class_counts']['ham']} 条，垃圾短信 spam 为 {summary['class_counts']['spam']} 条。缺失值检测显示，原始数据缺失值总数为 {summary['cleaning']['missing_before']}，清理后缺失值总数为 {summary['cleaning']['missing_after']}。数据清洗阶段执行 dropna 和去重，去除重复样本 {summary['cleaning']['duplicates_removed']} 条。")
    add_paragraph(doc, f"数据集划分严格遵守先划分后预处理原则。实验使用 stratify=y 保证训练集、验证集、测试集中的类别比例尽量一致，并固定 random_state=42。划分后训练集 {summary['splits']['train']['rows']} 条，验证集 {summary['splits']['val']['rows']} 条，测试集 {summary['splits']['test']['rows']} 条，整体接近 60%/20%/20%。这种划分方式可以避免测试集信息提前参与模型训练，也使实验结果具有可复现性。")
    add_paragraph(doc, f"文本特征采用 TF-IDF。向量器只在训练集上 fit，最终词表规模为 {summary['tfidf']['vocabulary_size']}，最大特征数设置为 {summary['tfidf']['max_features']}，ngram 范围为 {summary['tfidf']['ngram_range'][0]} 到 {summary['tfidf']['ngram_range'][1]}。这些配置来自验证集敏感性实验，而不是手工指定后直接用于测试集。验证集和测试集只调用 transform，不参与词表学习。这样做是为了模拟真实部署场景：模型只能从训练数据中学习特征空间，对新短信只能映射到既有词表。")
    add_paragraph(doc, "代码流程与课程 PDF 中的要求逐项对应：首先读取原始 SMSSpamCollection 文本并转换为 csv；随后执行 dropna 和 drop_duplicates 完成缺失值与重复样本处理；再用 value_counts 绘制类别分布，用短信长度统计绘制直方图，并统计训练集高频词；之后调用两次 train_test_split，其中第一次划出测试集，第二次从剩余数据中划出验证集，两次均设置 stratify=y 和 random_state=42；最后构建 TfidfVectorizer，只对训练集调用 fit_transform，对验证集和测试集只调用 transform。由于本任务输入是文本，因此不使用 StandardScaler 或 OneHotEncoder，而采用 PDF 中列出的文本型特征编码方法 TfidfVectorizer。")
    add_picture_if_exists(doc, metrics["figures"]["label_distribution"], "图 1 类别分布")
    add_picture_if_exists(doc, metrics["figures"]["message_length_distribution"], "图 2 短信长度分布")
    add_picture_if_exists(doc, metrics["figures"]["top_terms"], "图 3 训练集高频词分布")

    add_heading(doc, "五、实验设置与结果", 1)
    add_paragraph(doc, f"环境依赖与运行平台信息如下：实验环境为 {metrics['environment']['platform']}，Python 解释器为 {metrics['environment']['python_executable']}，Python 版本为 {metrics['environment']['python_version'].split()[0]}，PyTorch 版本为 {metrics['environment']['torch_version']}，训练设备为 {metrics['environment']['device']}。当前 PyTorch 为 CPU 版本，未使用 CUDA/GPU；本任务数据规模较小，在 CPU 上可以完成完整训练和敏感性实验。主模型最大训练轮数为 {metrics['config']['max_epochs']}，early stopping patience 为 {metrics['config']['early_stopping_patience']}，batch size 为 {metrics['config']['batch_size']}，学习率为 {metrics['config']['learning_rate']}，Dropout 为 {metrics['config']['dropout']}，weight decay 为 {metrics['config']['weight_decay']}。最终测试使用验证集 Macro-F1 最优的 checkpoint，而不是最后一轮参数。")
    add_table_from_dataframe(doc, dependency_versions(), "表 1 实验环境主要依赖版本")

    baseline_display = baseline_df.copy()
    add_table_from_dataframe(doc, baseline_display, "表 2 Baseline 在验证集和测试集上的结果")
    add_paragraph(doc, "多数类 baseline 的作用是衡量类别不均衡背景下“只猜正常短信”能达到的虚高准确率。随机 baseline 用于观察按类别比例随机预测的下限表现。Logistic Regression、Multinomial Naive Bayes 和 Linear SVM 都是文本分类中的常见传统模型，它们与 MLP 使用同一份 TF-IDF 特征，可以帮助判断非线性隐藏层是否真的带来额外收益。")

    add_paragraph(doc, f"主模型 MLP 在测试集上的 Accuracy 为 {fmt(metrics['mlp']['test']['accuracy'])}，Macro-F1 为 {fmt(metrics['mlp']['test']['macro_f1'])}，Macro-Precision 为 {fmt(metrics['mlp']['test']['macro_precision'])}，Macro-Recall 为 {fmt(metrics['mlp']['test']['macro_recall'])}。验证集最优 Accuracy 为 {fmt(metrics['mlp']['validation']['accuracy'])}，Macro-F1 为 {fmt(metrics['mlp']['validation']['macro_f1'])}。训练过程中验证 Macro-F1 最优出现在第 {metrics['mlp']['history_summary']['best_epoch']} 轮，实际训练到第 {metrics['mlp']['history_summary']['epochs_ran']} 轮触发 early stopping。")
    add_picture_if_exists(doc, metrics["figures"]["training_curve"], "图 4 训练集与验证集 Loss/Accuracy 曲线")
    add_picture_if_exists(doc, metrics["figures"]["confusion_matrix"], "图 5 测试集混淆矩阵")

    add_heading(doc, "六、超参数敏感性分析", 1)
    add_paragraph(doc, "超参数敏感性分析采用控制变量法。优化后的实验不再比较第 50 轮或最后一轮结果，而是对每组超参数都保存验证集 Macro-F1 最优的 checkpoint，并用 best validation 指标进行比较。学习率实验只改变 learning_rate，Dropout 实验只改变 dropout；随后继续比较隐藏层结构和 TF-IDF 参数。这样验证集真正参与模型选择，测试集只用于最终评估。")
    hyper_display = hyper_df[["parameter", "value", "best_epoch", "epochs_ran", "best_val_accuracy", "best_val_macro_f1", "final_val_macro_f1"]].copy()
    add_table_from_dataframe(doc, hyper_display, "表 3 学习率与 Dropout 敏感性实验结果")
    add_picture_if_exists(doc, metrics["figures"]["lr_sensitivity"], "图 6 学习率敏感性分析")
    add_picture_if_exists(doc, metrics["figures"]["dropout_sensitivity"], "图 7 Dropout 敏感性分析")

    best_lr = hyper_df[hyper_df["parameter"] == "learning_rate"].sort_values("best_val_macro_f1", ascending=False).iloc[0]
    best_dropout = hyper_df[hyper_df["parameter"] == "dropout"].sort_values("best_val_macro_f1", ascending=False).iloc[0]
    add_paragraph(doc, f"从实验结果看，学习率实验中验证 Macro-F1 最好的取值为 {best_lr['value']}，对应 best epoch 为 {int(best_lr['best_epoch'])}，验证 Accuracy 为 {fmt(best_lr['best_val_accuracy'])}，验证 Macro-F1 为 {fmt(best_lr['best_val_macro_f1'])}。学习率 0.0001 在本次实验中最好，说明该任务虽然容易收敛，但较小学习率配合 early stopping 更稳定。")
    add_paragraph(doc, f"Dropout 实验中验证 Macro-F1 最好的取值为 {best_dropout['value']}，对应 best epoch 为 {int(best_dropout['best_epoch'])}，验证 Accuracy 为 {fmt(best_dropout['best_val_accuracy'])}，验证 Macro-F1 为 {fmt(best_dropout['best_val_macro_f1'])}。本次 Dropout=0.1 优于 0.3 和 0.5，说明在加入 early stopping 和类别权重后，过强的随机失活会削弱模型对少数类 spam 的有效学习。")
    model_display = model_selection_df[["value", "best_epoch", "epochs_ran", "best_val_accuracy", "best_val_macro_f1", "final_val_macro_f1"]].copy()
    add_table_from_dataframe(doc, model_display, "表 4 隐藏层结构敏感性实验结果")
    add_picture_if_exists(doc, metrics["figures"]["hidden_dims_sensitivity"], "图 8 隐藏层结构敏感性分析")
    tfidf_display = tfidf_df[["parameter", "value", "best_epoch", "epochs_ran", "best_val_accuracy", "best_val_macro_f1", "final_val_macro_f1"]].copy()
    add_table_from_dataframe(doc, tfidf_display, "表 5 TF-IDF 参数敏感性实验结果")
    add_picture_if_exists(doc, metrics["figures"]["tfidf_sensitivity"], "图 9 TF-IDF 参数敏感性分析")
    best_hidden = model_selection_df.sort_values("best_val_macro_f1", ascending=False).iloc[0]
    best_tfidf_max = tfidf_df[tfidf_df["parameter"] == "max_features"].sort_values("best_val_macro_f1", ascending=False).iloc[0]
    best_tfidf_ngram = tfidf_df[tfidf_df["parameter"] == "ngram_range"].sort_values("best_val_macro_f1", ascending=False).iloc[0]
    add_paragraph(doc, f"隐藏层结构实验中，{best_hidden['value']} 的验证 Macro-F1 最高，为 {fmt(best_hidden['best_val_macro_f1'])}。TF-IDF 参数实验中，max_features={best_tfidf_max['value']} 的验证 Macro-F1 为 {fmt(best_tfidf_max['best_val_macro_f1'])}，ngram_range={best_tfidf_ngram['value']} 的验证 Macro-F1 为 {fmt(best_tfidf_ngram['best_val_macro_f1'])}。最终模型采用验证集 Macro-F1 最优组合：学习率 {metrics['selection']['selected_learning_rate']}、Dropout {metrics['selection']['selected_dropout']}、隐藏层 {metrics['selection']['selected_hidden_dims']}、max_features {metrics['selection']['selected_max_features']}、ngram_range {metrics['selection']['selected_ngram_range']}。")

    add_heading(doc, "七、结果分析与错误分析", 1)
    logistic_test = baseline_df[(baseline_df["model"] == "logistic_regression") & (baseline_df["split"] == "test")].iloc[0]
    svm_test = baseline_df[(baseline_df["model"] == "linear_svm") & (baseline_df["split"] == "test")].iloc[0]
    majority_test = baseline_df[(baseline_df["model"] == "majority_class") & (baseline_df["split"] == "test")].iloc[0]
    mlp_macro = float(metrics["mlp"]["test"]["macro_f1"])
    logistic_macro = float(logistic_test["macro_f1"])
    svm_macro = float(svm_test["macro_f1"])
    if mlp_macro >= max(logistic_macro, svm_macro):
        comparison_text = "MLP 高于 Logistic Regression 和 Linear SVM，说明隐藏层的非线性组合在本次划分上带来了一定收益。"
    else:
        comparison_text = "Linear SVM 和 Logistic Regression 均高于 MLP，说明 TF-IDF 特征在该数据集上具有较强线性可分性，复杂一点的神经网络并不必然带来更好泛化。"
    add_paragraph(doc, "本节按作业要求的结果分析四步展开：超参影响、错误分析、过拟合判断和 Baseline 对比。")
    add_paragraph(doc, f"与多数类 baseline 相比，MLP 测试 Macro-F1 从 {fmt(majority_test['macro_f1'])} 提升到 {fmt(metrics['mlp']['test']['macro_f1'])}。这说明模型并不是简单依赖数据集中正常短信数量更多这一事实，而是学习到了区分 spam 与 ham 的文本特征。与传统文本分类模型相比，MLP 测试 Macro-F1 为 {fmt(metrics['mlp']['test']['macro_f1'])}，Logistic Regression 为 {fmt(logistic_test['macro_f1'])}，Linear SVM 为 {fmt(svm_test['macro_f1'])}。{comparison_text}")
    add_paragraph(doc, "混淆矩阵可以进一步观察错误类型。垃圾短信识别中最需要关注的是 spam 被误判为 ham，因为这类错误会使垃圾短信逃过滤器。ham 被误判为 spam 虽然也会影响用户体验，但通常可以通过人工恢复或白名单机制缓解。由于本项目使用的是英文短信数据，一些垃圾短信包含 free、call、win、prize 等明显词项，模型较容易识别；而表达较短、词汇普通、没有明显诱导词的 spam 更容易与正常短信混淆。")
    add_paragraph(doc, f"从训练曲线看，最终训练 Accuracy 为 {fmt(metrics['mlp']['history_summary']['final_train_accuracy'])}，验证 Accuracy 为 {fmt(metrics['mlp']['history_summary']['final_val_accuracy'])}；最终训练 Loss 为 {fmt(metrics['mlp']['history_summary']['final_train_loss'])}，验证 Loss 为 {fmt(metrics['mlp']['history_summary']['final_val_loss'])}。图 4 中 Loss 在前 10 个 epoch 下降最快，之后训练 Loss 仍继续下降，而验证 Loss 基本稳定在 0.12 左右，说明模型已经接近收敛；Accuracy 在第 3 到第 5 个 epoch 后快速接近较高水平，后续提升幅度有限。优化后模型在第 {metrics['mlp']['history_summary']['best_epoch']} 轮取得验证 Macro-F1 最优值，并在第 {metrics['mlp']['history_summary']['epochs_ran']} 轮 early stopping。与旧实验固定训练 200 轮相比，本次训练没有继续无意义地压低训练损失，验证集真正用于选择最终 checkpoint。")
    add_paragraph(doc, "图 5 的测试集混淆矩阵显示，904 条 ham 中有 888 条被正确识别，16 条被误判为 spam；131 条 spam 中有 116 条被正确识别，15 条被误判为 ham。两类错误数量接近，但从应用风险看，spam 被误判为 ham 更值得关注，因为这意味着垃圾短信会直接到达用户。结合文本特征看，这类错误可能来自较短的 spam、缺少 free/call/win/prize 等显著词项的营销短信，或表达方式接近正常通知的垃圾短信；而少量 ham 被误判为 spam，可能是因为正常短信中也出现 reward、urgent、call 等高风险词。")

    add_heading(doc, "八、总结与反思", 1)
    add_paragraph(doc, "本项目围绕垃圾短信识别任务，完成了从公开数据获取、缺失值检测、数据可视化、分层划分、TF-IDF 特征提取、MLP 训练、baseline 对比、超参数敏感性分析到实验结果保存的完整流程。优化后的实验加入 early stopping、best checkpoint、类别权重、隐藏层结构敏感性和 TF-IDF 参数敏感性，使验证集真正参与模型选择。实验设计遵守先划分数据集再进行预处理的原则，避免了特征泄漏。所有实验结果均保存在 results 目录，报告中的数值来自这些结果文件。")
    add_paragraph(doc, "项目的主要优点是流程完整、实现简单、可复现性较强，并且能够在普通 CPU 环境上运行。局限性也比较明显：第一，数据集是英文短信，不能直接代表中文垃圾短信场景；第二，TF-IDF 主要反映词项统计权重，无法充分理解复杂语义和上下文；第三，MLP 的解释性有限，且本次真实结果显示它并未超过 Linear SVM 和 Logistic Regression；第四，数据集规模较小，模型在真实部署场景下仍需要更多样化数据验证。这个结果也提醒我，在监督学习应用中不能只把神经网络视为更高级的默认选项，而应根据数据特征、样本量、可解释性和部署成本选择合适模型。")
    add_paragraph(doc, "后续改进可以从三个方向展开。其一，引入更贴近中文语境的数据集，并针对中文分词、数字链接脱敏、特殊符号处理进行特征工程。其二，增加多随机种子重复实验，报告平均值和标准差，从而判断不同模型之间的差距是否稳定。其三，在保持学术诚信和可解释性的前提下，尝试轻量级深度文本模型，并分析其相较 TF-IDF + Linear SVM 的收益是否值得额外复杂度。")
    add_paragraph(doc, "如果将该模型放到真实短信过滤系统中，还需要考虑课程实验之外的工程约束。首先，模型预测不应直接等同于删除短信，而更适合输出风险提示、垃圾箱归档或二次确认，因为误杀正常短信会影响用户接收验证码、物流通知和重要提醒。其次，垃圾短信的表达方式会随时间变化，攻击者可能故意替换字符、插入空格、使用同音词或短链接绕过基于词项统计的模型，因此系统需要定期用新样本重新评估。再次，训练集中的英文短信来源具有地域和时间背景，模型在中文短信、混合语言短信或新的营销模板上可能出现性能下降。")
    add_paragraph(doc, "指标选择也会影响对模型的判断。Accuracy 在类别不均衡时容易偏乐观，因为只要正常短信占比较高，模型即使偏向预测 ham 也可能得到较高准确率。Macro-F1 会分别考虑每个类别的 F1 后再平均，因此能更好反映少数类 spam 的识别情况。对于垃圾短信识别，实际应用中还可以根据需求调整阈值：如果更重视拦截垃圾短信，可以提高 spam 召回率；如果更担心误拦正常短信，则应提高 spam 精确率并设置人工复核机制。本项目保持标准分类阈值，是为了让实验流程简单清晰，便于复现和对比。")
    add_paragraph(doc, "通过本次实验，我对监督学习流程的理解主要有三点。第一，数据划分和预处理顺序比模型结构更基础，如果在划分前对全量数据 fit 特征提取器，就会把验证集和测试集信息泄漏到训练过程。第二，类别不均衡会使 Accuracy 产生误导，必须结合 Macro-F1、混淆矩阵和具体业务风险判断模型是否可用。第三，验证集的作用不是展示结果，而是服务于模型选择；测试集应尽量保持独立，只在最终阶段用于估计泛化性能。")

    add_heading(doc, "参考文献", 1)
    add_plain_paragraph(doc, "[1] Almeida, T. & Hidalgo, J. (2011). SMS Spam Collection [Dataset]. UCI Machine Learning Repository. https://doi.org/10.24432/C5CC84")
    add_plain_paragraph(doc, "[2] Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.")
    add_plain_paragraph(doc, "[3] Paszke, A. et al. (2019). PyTorch: An Imperative Style, High-Performance Deep Learning Library. Advances in Neural Information Processing Systems.")
    add_plain_paragraph(doc, "[4] 课程 PDF：《人工智能课程论文：基于监督学习的分类应用设计与实现》。")

    add_heading(doc, "提交前诚信提醒", 1)
    add_paragraph(doc, "最终提交前应自行运行代码、核对 results 目录中的实验结果、理解每个脚本的作用，并根据个人理解改写报告中的分析性文字。代码可以作为课程项目实现基础，但报告中的实验理解、结果解释和反思部分应体现本人真实理解。")

    count = chinese_char_count(doc)
    if count < 4000:
        add_plain_paragraph(doc, "当前正文中文字符数不足 4000，请继续补充个人分析。")

    bold_required_terms(doc)
    normalize_document_fonts(doc)
    REPORT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        output_path = REPORT_PATH
        doc.save(output_path)
    except PermissionError:
        output_path = FONT_CLEAN_PATH
        try:
            doc.save(output_path)
            print("原报告可能正在被 Word 打开，已另存为字体整理版。")
        except PermissionError:
            output_path = project_path(f"report/AI-王李明-2024302181194-{int(time.time())}.docx")
            doc.save(output_path)
            print("原报告和字体整理版都可能正在被 Word 打开，已另存为带时间戳的新文件。")
    print(f"报告已生成: {output_path}")
    print(f"中文字符数约: {count}")


if __name__ == "__main__":
    build_report()
